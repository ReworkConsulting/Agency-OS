"""
export_to_docx.py

Converts a markdown file to a formatted .docx document.
Usage: python3 tools/export_to_docx.py <input.md> <output.docx>
"""

import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_heading_style(paragraph, level, text):
    """Apply heading formatting."""
    run = paragraph.add_run(text)
    if level == 1:
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    elif level == 2:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x16, 0x21, 0x3E)
    elif level == 3:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0F, 0x3C, 0x78)
    elif level == 4:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return paragraph


def add_table_from_md(doc, lines, start_idx):
    """Parse a markdown table and add it to the doc. Returns the index after the table."""
    table_lines = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith('|'):
        table_lines.append(lines[i].strip())
        i += 1

    if len(table_lines) < 2:
        return i

    # Parse header
    header_cells = [c.strip() for c in table_lines[0].split('|') if c.strip()]
    # Skip separator row (line with dashes)
    data_rows = []
    for row_line in table_lines[2:]:
        cells = [c.strip() for c in row_line.split('|') if c.strip()]
        if cells:
            data_rows.append(cells)

    if not header_cells:
        return i

    col_count = len(header_cells)
    table = doc.add_table(rows=1 + len(data_rows), cols=col_count)
    table.style = 'Table Grid'

    # Header row
    hdr_row = table.rows[0]
    for j, cell_text in enumerate(header_cells[:col_count]):
        cell = hdr_row.cells[j]
        cell.text = cell_text
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10)
        # Light blue header background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D6E4F0')
        tcPr.append(shd)

    # Data rows
    for r, row_data in enumerate(data_rows):
        row = table.rows[r + 1]
        for j in range(col_count):
            cell_text = row_data[j] if j < len(row_data) else ''
            cell = row.cells[j]
            cell.text = apply_inline_formatting_plain(cell_text)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    return i


def apply_inline_formatting_plain(text):
    """Strip markdown bold/italic markers and return plain text."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text


def add_formatted_paragraph(doc, text, style='Normal'):
    """Add a paragraph with inline bold/italic formatting preserved."""
    para = doc.add_paragraph(style=style)
    para.paragraph_format.space_after = Pt(6)

    # Split on bold (**text**) and italic (*text*) markers
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        else:
            para.add_run(part)

    return para


def add_blockquote(doc, text):
    """Add a styled blockquote paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.4)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)

    # Add left border via XML
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '10')
    left.set(qn('w:color'), '4A90D9')
    pBdr.append(left)
    pPr.append(pBdr)

    run = para.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return para


def convert_md_to_docx(md_path, docx_path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Default paragraph font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    blockquote_buffer = []

    def flush_blockquote():
        nonlocal blockquote_buffer
        if blockquote_buffer:
            text = ' '.join(blockquote_buffer).strip()
            add_blockquote(doc, text)
            blockquote_buffer = []

    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # Blockquote
        if stripped.startswith('> '):
            blockquote_buffer.append(stripped[2:])
            i += 1
            continue
        else:
            flush_blockquote()

        # Horizontal rule
        if stripped in ('---', '***', '___'):
            para = doc.add_paragraph()
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,4})\s+(.+)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            para = doc.add_paragraph()
            if level > 1:
                para.paragraph_format.space_before = Pt(14)
            set_heading_style(para, level, text)
            i += 1
            continue

        # Table
        if stripped.startswith('|'):
            i = add_table_from_md(doc, lines, i)
            continue

        # Bullet list
        bullet_match = re.match(r'^[-*]\s+(.+)', stripped)
        if bullet_match:
            content = bullet_match.group(1)
            para = doc.add_paragraph(style='List Bullet')
            para.paragraph_format.space_after = Pt(3)
            parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = para.add_run(part[1:-1])
                    run.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = para.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                else:
                    para.add_run(part)
            i += 1
            continue

        # Numbered list
        num_match = re.match(r'^\d+\.\s+(.+)', stripped)
        if num_match:
            content = num_match.group(1)
            para = doc.add_paragraph(style='List Number')
            para.paragraph_format.space_after = Pt(3)
            parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = para.add_run(part[1:-1])
                    run.italic = True
                else:
                    para.add_run(part)
            i += 1
            continue

        # Checkbox list item
        checkbox_match = re.match(r'^- \[[ x]\] (.+)', stripped)
        if checkbox_match:
            content = checkbox_match.group(1)
            checked = '☑' if '- [x]' in stripped else '☐'
            para = doc.add_paragraph(style='List Bullet')
            para.paragraph_format.space_after = Pt(3)
            para.add_run(f'{checked} {content}')
            i += 1
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Normal paragraph
        add_formatted_paragraph(doc, stripped)
        i += 1

    flush_blockquote()

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("Usage: python3 export_to_docx.py <input.md> <output.docx>")
        sys.exit(1)
    convert_md_to_docx(sys.argv[1], sys.argv[2])
